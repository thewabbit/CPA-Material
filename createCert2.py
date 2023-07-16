import datetime
import multiprocessing.dummy
import os
import random
import re
import shutil
import tempfile
import time
import tkinter as tk
import uuid
from tkinter import Button, Label, filedialog, scrolledtext
import json

import docx
import docx2pdf
import openpyxl
import PyPDF2
import tkcalendar

class templates():
    def __init__(self):
        self.certificateTemplate = None
        self.speakerTemplate = None
        self.weighinTemplate = None
        self.gearcheckTemplate = None
        self.lifterDataInput = None
        self.manualscorecardTemplate = None
        self.olTemplate = None
        self.lifterJSONTemplate = None
        self.eventName = None

class CPA_Cert_Generator (tk.Tk):
    def __init__(self):
        """Initialize the GUI application with buttons, labels, and a text box for logging."""

        self.input = templates()

        super().__init__()

        self.lifterData =[]
        self.GUID = ""
        # self.temp = "c:\\temp\\"
        self.temp = ""
        self.lots = {}
        self.days = 0
        self.eventDate = ""

        # Set the initial values of the template and lifterData attributes to None.
        self.templateNames = {
            "certificateTemplate":"3LIFT CERT TEMPLATE.docx",
            "speakerTemplate":"SPEAKER TEMPLATE.docx",
            "weighinTemplate":"WEIGH IN TEMPLATE.docx",
            "gearcheckTemplate":"GEAR CHECK TEMPLATE.docx",
            "manualscorecardTemplate":"MANUAL SCORESHEET TEMPLATE.docx",
            "olTemplate":"TEMPLATE.openlifter",
            "lifterJSONTemplate":"LIFTER TEMPLATE.json"
        }



        self.certificatePDFs = []
        self.speakerPDFs = []

        # Set the size, background color, and title of the window.
        self.geometry('800x720')
        self.configure(background='#FAEBD7')
        self.title('CPA Certificate Generator')

        # Set the fonts for buttons, labels, and file path displays.
        self.buttonFont = ('courier', 12, 'normal')
        self.labelFont = ('courier', 12, 'bold')
        self.pathFont = ('courier', 12, 'normal')
        self.logFont = ('courier', 10, 'normal')

        # Set up the logging text box and configure it with tags for different message types.
        self.logBox = scrolledtext.ScrolledText(self, wrap=tk.WORD, width=85, height=10, font=self.logFont)
        self.logBox.place(x=40, y=520)
        self.logBox.insert(tk.INSERT, 'Please select files \n')
        self.logBox.configure(state='disabled')
        self.logBox.tag_configure('error', foreground='red')
        self.logBox.tag_configure('warning', foreground='orange')


        def create_button(parent, name, text, x, y, command):
            button = Button(parent, name=name, text=text, bg='#FAEBD7', font=self.buttonFont, command=command)
            button.place(x=x, y=y)
            return button

        def create_label(parent, name, text, x, y):
            label = Label(parent, name=name, text=text, bg='#FAEBD7', font=self.pathFont)
            label.place(x=x, y=y)
            return label
        
        self.eventNameField = tk.Entry(self, font=self.labelFont)
        self.eventNameField.place(x=40,y=50)
        self.eventNameLabel = create_label(self, 'eventNameLabel', 'Event Name', 40, 20)

        self.eventDateField = tkcalendar.DateEntry(self, date_pattern='yyyy-mm-dd')
        self.eventDateField.place(x=40,y=110)
        self.eventDateLabel = create_label(self, 'eventDateLabel', 'Event Starting Date', 40, 80)

        self.OLTemplateButton = create_button(self, 'olTemplateButton', 'Openlifter Template', 40, 180, self.selectOLTemplate)
        self.selectedOlTemplate = create_label(self, 'selectedOlTemplate', '<please select file>', 40, 220)

        self.LifterJSONTemplateButton = create_button(self, 'lifterJSONButton', 'Lifter JSON Template', 40, 260, self.selectOLTemplate)
        self.selectedLifterJSONTemplate = create_label(self, 'selectedLifterJSONTemplate', '<please select file>', 40, 300)

        self.LifterDataButton = create_button(self, 'lifterDataButton', 'Lifter Data', 350, 50, self.selectLifterData)
        self.selectedLifterData = create_label(self, 'selectedLifterData', '<please select file>', 350, 90)

        self.certificateTemplateButton = create_button(self, 'certificateTemplateButton', 'Certificate Template', 350, 180, self.selectCertificateTemplate)
        self.selectedCertificateTemplate = create_label(self, 'selectedCertificateTemplate', '<please select file>', 350, 220)

        self.speakerTemplateButton = create_button(self, 'speakerTemplateButton', 'Speaker Card Template', 350, 260, self.selectSpeakerTemplate)
        self.selectedSpeakerTemplate = create_label(self, 'selectedSpeakerTemplate', '<please select file>', 350, 300)

        self.weighinTemplateButton = create_button(self, 'weighinTemplateButton', 'Weigh In Template', 350, 340, self.selectWeighinTemplate)
        self.selectedWeighinTemplate = create_label(self, 'selectedWeighinTemplate', '<please select file>', 350, 380)

        self.gearcheckTemplateButton = create_button(self, 'gearcheckTemplateButton', 'Gear Check Template', 350, 420, self.selectGearcheckTemplate)
        self.selectedGearcheckTemplate = create_label(self, 'selectedGearcheckTemplate', '<please select file>', 350, 470)

        self.manualscorecardTemplateButton = create_button(self, 'manualscorecardTemplateButton', 'Manual Score Card Template', 40, 340, self.selectManualScoreCardTemplate)
        self.selectedManualscorecardTemplate = create_label(self, 'selectedManualscorecardTemplate', '<please select file>', 40, 380)

        self.runButton = Button(self, text='Create files', bg='#FAEBD7', font=self.buttonFont, command=self.run)
        self.runButton.place(x=40, y=480)

        self.templateScan()

    def templateScan(self):
        for k,v in self.templateNames.items():
            path = os.path.join(os.getcwd(),"Templates",v)
            if os.path.exists(path):
                setattr(self.input,k,path)
                v1 = f"selected{k[0].upper() + k[1:]}"
                getattr(self, v1).configure(text = v)

    def selectOLTemplate(self):
        # Open a file dialog box to allow the user to select the lifter data Excel file
        self.input.olTemplate = filedialog.askopenfilename(filetypes=[('OL Template', '*.openlifter')])
        
        # Update the GUI to show the selected file name
        self.selectedOlTemplate.configure(text = os.path.split(self.input.olTemplate)[1])
        
        # Log the selected file path in the log box
        self.log(f'selected input data:\n{self.input.olTemplate}')

    def selectLifterJSONTemplate(self):
        # Open a file dialog box to allow the user to select the lifter data Excel file
        self.input.lifterJSONTemplate = filedialog.askopenfilename(filetypes=[('OL Template', '*.openlifter')])
        
        # Update the GUI to show the selected file name
        self.selectedLifterJSONTemplate.configure(text = os.path.split(self.input.lifterJSONTemplate)[1])
        
        # Log the selected file path in the log box
        self.log(f'selected input data:\n{self.input.lifterJSONTemplate}')
    
    def selectLifterData(self):
        # Open a file dialog box to allow the user to select the lifter data Excel file
        self.input.lifterDataInput = filedialog.askopenfilename(filetypes=[('Lifter data', '*.xlsx')])
        
        # Update the GUI to show the selected file name
        self.selectedLifterData.configure(text = os.path.split(self.input.lifterDataInput)[1])
        
        # Log the selected file path in the log box
        self.log(f'selected input data:\n{self.input.lifterDataInput}')

    def selectCertificateTemplate(self):
        # Open a file dialog box to allow the user to select the certificate template Word file
        self.input.certificateTemplate = filedialog.askopenfilename(filetypes=[('Template file', '*.docx')])
        
        # Update the GUI to show the selected file name
        self.selectedCertificateTemplate.configure(text = os.path.split(self.input.certificateTemplate)[1])
        
        # Log the selected file path in the log box
        self.log(f'selected template:\n{self.input.certificateTemplate}')

    def selectSpeakerTemplate(self):
        # Open a file dialog box to allow the user to select the certificate template Word file
        self.input.speakerTemplate = filedialog.askopenfilename(filetypes=[('Template file', '*.docx')])
        
        # Update the GUI to show the selected file name
        self.selectedSpeakerTemplate.configure(text = os.path.split(self.input.speakerTemplate)[1])
        
        # Log the selected file path in the log box
        self.log(f'selected template:\n{self.input.speakerTemplate}')

    def selectWeighinTemplate(self):
        # Open a file dialog box to allow the user to select the certificate template Word file
        self.input.weighinTemplate = filedialog.askopenfilename(filetypes=[('Template file', '*.docx')])
        
        # Update the GUI to show the selected file name
        self.selectedWeighinTemplate.configure(text = os.path.split(self.input.weighinTemplate)[1])
        
        # Log the selected file path in the log box
        self.log(f'selected template:\n{self.input.weighinTemplate}')

    def selectGearcheckTemplate(self):
        # Open a file dialog box to allow the user to select the certificate template Word file
        self.input.gearcheckTemplate = filedialog.askopenfilename(filetypes=[('Template file', '*.docx')])
        
        # Update the GUI to show the selected file name
        self.selectedGearcheckTemplate.configure(text = os.path.split(self.input.gearcheckTemplate)[1])
        
        # Log the selected file path in the log box
        self.log(f'selected template:\n{self.input.gearcheckTemplate}')

    def selectManualScoreCardTemplate(self):
        # Open a file dialog box to allow the user to select the certificate template Word file
        self.input.manualscorecardTemplate = filedialog.askopenfilename(filetypes=[('Template file', '*.docx')])
        
        # Update the GUI to show the selected file name
        self.selectedManualscorecardTemplate.configure(text = os.path.split(self.input.manualscorecardTemplate)[1])
        
        # Log the selected file path in the log box
        self.log(f'selected template:\n{self.input.manualscorecardTemplate}')


    def log(self, message, level='info'):
        """
        Logs a message to the GUI log box.

        Args:
            message (str): The message to log.
            level (str, optional): The level of the message (default 'info').

        Returns:
            None
        """
        # Enables the logBox for writing
        self.logBox.configure(state='normal')
        # Writes the message to the logBox, using the given level
        self.logBox.insert(tk.INSERT, message + "\n", level)
        # Disables the logBox after writing
        self.logBox.configure(state='disabled')
        # Scrolls to the end of the logBox
        self.logBox.see("end")
        # Updates the GUI
        self.update()

    def proccessData(self):
        self.log('compiling lifter input data')

        self.input.eventName = self.eventNameField.get()

        # Load the Excel workbook and select the active worksheet
        wb = openpyxl.load_workbook(self.input.lifterDataInput)
        ws = wb.active

        ID = 1
        # Iterate through each row of data in the worksheet
        for row in ws.iter_rows(min_row=2, values_only=True):
            # Check if the first column of the row is not None
            if row[0] != None:
                day = row[0]
                flight = row[1]
                __rand = str(random.randint(0, 100000))

                if day in self.lots:
                    if flight not in self.lots[day]:
                        self.lots[day][flight] = []
                else:
                    self.lots[day] = {}
                    if flight not in self.lots[day]:
                        self.lots[day][flight] = []

                self.lots[day][flight].append({"id":ID,"rand":__rand})

                lifterDataTemplate = {
                    "Day": day,
                    "Flight": flight,
                    "Age": re.sub(r'\([^)]*\)', '', row[2]),
                    "Weight": re.sub(r'^.*\s-\s', '', row[3]),
                    "First name": row[4],
                    "Last name": row[5],
                    "Gender": row[6],
                    "DOB": row[7],
                    "Nation": "New Zealand",
                    "Association": row[8],
                    "Raw": row[9],
                    "Instagram": row[10],
                    "Notes": row[11],
                    "ID": ID,
                    "Lot": "0"
                }

                self.lifterData.append(lifterDataTemplate)
                ID += 1

        #set lot numbers
        for day, flight in self.lots.items():
            print(day)
            for flight, v in flight.items():
                print(flight)
                lotsSorted = sorted(v, key=lambda x: x['rand'])
                for i in range(0, len(lotsSorted)):
                    self.lifterData[lotsSorted[i]["id"]-1]["Lot"]=str(i+1)
        
        #get number of days
        self.days = max(list(set(item["Day"] for item in self.lifterData)))

    def findReplaceParagraph(self, doc, data, revert=False):
        # If revert is True, swap find and replace arguments
        if revert:
            data = {v: k for k, v in data.items()}

        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, value)

    def createGearCheck(self):
        allPDFs = []
        self.log("Creating gear check sheets")
        for i in range(0,self.days):
            i += 1
            curDay = [d for d in self.lifterData if d['Day'] == i]

            self.log(f"Gear check day {i}")

            doc = docx.Document(self.input.gearcheckTemplate)
            table = doc.tables[0]
            for j in sorted(curDay, key=lambda x: (x['Flight'],int(x['Lot']))):
                rc = table.add_row().cells
                rc[0].text = f"{j['First name']} {j['Last name']}"  
        
            
            self.findReplaceParagraph(doc,dict(zzEVENTzz=self.input.eventName,zzSESSIONzz=str(i)))

            time.sleep(0.5)
            path = os.path.join(self.temp, f'gear_check_{str(i)}.docx')
            doc.save(path)
            time.sleep(0.5)

            docx2pdf.convert(path)

            # Add the path of the generated PDF to the list
            allPDFs.append(os.path.join(self.temp, f'gear_check_{str(i)}.pdf'))
        
        

        # Create a new PDF merger object
        merger = PyPDF2.PdfMerger()

        # Iterate through each PDF file path in the list and add it to the merger
        for pdf in allPDFs:
            merger.append(pdf)

        # Write the merged PDF to a file
        merger.write('gear check.pdf')
        self.log("Gear check created")

    def createWeighIn(self):
        allPDFs = []
        self.log("Creating weigh ins")
        for i in range(0,self.days):
            i += 1
            curDay = [d for d in self.lifterData if d['Day'] == i]

            self.log(f"Weigh ins day {i}")

            doc = docx.Document(self.input.weighinTemplate)
            table = doc.tables[0]
            for j in sorted(curDay, key=lambda x: (x['Flight'],int(x['Lot']))):
                rc = table.add_row().cells
                rc[0].text = str(j['Day'])
                rc[1].text = j["Flight"]
                rc[2].text = f"{j['First name']} {j['Last name']}"
                rc[3].text = j['Lot']            
            
            time.sleep(0.5)
            path = os.path.join(self.temp, f'weigh_in_day_{str(i)}.docx')
            doc.save(path)
            time.sleep(0.5)

            docx2pdf.convert(path)

            # Add the path of the generated PDF to the list
            allPDFs.append(os.path.join(self.temp, f'weigh_in_day_{str(i)}.pdf'))

        
        
        # Create a new PDF merger object
        merger = PyPDF2.PdfMerger()

        # Iterate through each PDF file path in the list and add it to the merger
        for pdf in allPDFs:
            merger.append(pdf)

        # Write the merged PDF to a file
        merger.write('weigh in.pdf')
        self.log("Weigh ins created")
        
    def runLifterSpecific(self):

        self.log(f"Generating certificates and speaker cards\nThis will take a couple of minutes")
        #pool1= multiprocessing.Pool(processes=multiprocessing.cpu_count())
        pool1= multiprocessing.Pool(1)

        pdfgen = PDFGenerator()

        args =  [(lifter, self.input.certificateTemplate, self.temp) for lifter in self.lifterData]
        resultsCert = pool1.map(pdfgen.createCetificates, args)
        args =  [(lifter, self.input.speakerTemplate, self.temp) for lifter in self.lifterData]
        resultsSpeaker = pool1.map(pdfgen.createSpeaker, args)

        # close the pool and wait for the work to finish
        pool1.close()
        pool1.join()

        #docx2pdf.convert all docx in temp dir to pdfs
        docx2pdf.convert(self.temp)

        merger1 = PyPDF2.PdfMerger()
        # Iterate through each PDF file path in the list and add it to the merger
        for pdf in resultsCert:
            merger1.append(pdf.replace('.docx','.pdf'))
        # Write the merged PDF to a file
        merger1.write('certificates.pdf')


        merger2 = PyPDF2.PdfMerger()
        # Iterate through each PDF file path in the list and add it to the merger
        for pdf in resultsSpeaker:
            merger2.append(pdf.replace('.docx','.pdf'))
        # Write the merged PDF to a file
        merger2.write('speaker cards.pdf')

        self.log(f"Finished generating certificates and speaker cards")

    def createManualScoreSheet(self):
            allPDFs = []
            self.log("Creating manual scoresheets")
            for i in range(0,self.days):
                i += 1
                curDay = [d for d in self.lifterData if d['Day'] == i]

                self.log(f"Scoresheet day {i}")

                doc = docx.Document(self.input.manualscorecardTemplate)
                table = doc.tables[0]
                for j in sorted(curDay, key=lambda x: (x['Flight'],int(x['Lot']))):
                    rc = table.add_row().cells
                    rc[0].text = f"{j['First name']} {j['Last name']}"  
                    rc[1].text = f"" #teamS
                    rc[2].text = f"{j['Age']}"
                    rc[3].text = f"{j['Lot']}"
                    rc[5].text = f"{j['Weight']}"
            
                
                self.findReplaceParagraph(doc,dict(zzEVENTzz=self.input.eventName,zzSESSIONzz=str(i)))

                time.sleep(0.5)
                path = os.path.join(self.temp, f'manual_scoresheet_{str(i)}.docx')
                doc.save(path)
                time.sleep(0.5)
                docx2pdf.convert(path)

                # Add the path of the generated PDF to the list
                allPDFs.append(os.path.join(self.temp, f'manual_scoresheet_{str(i)}.pdf'))
            
            

            # Create a new PDF merger object
            merger = PyPDF2.PdfMerger()

            # Iterate through each PDF file path in the list and add it to the merger
            for pdf in allPDFs:
                merger.append(pdf)

            # Write the merged PDF to a file
            merger.write('manual scoresheet.pdf')
            self.log("Scoresheet created")

    def cleanUp(self):
        #seems to have an issues remove locked files

        # time.sleep(5)
        # shutil.rmtree(os.path.join(self.temp, self.GUID))
        # self.log("Cleaning up")
        pass


    def createOLData(self):
        self.eventDate = self.eventDateField.get_date()

        for i in range(0,self.days):
            i += 1
            id = 0
            outputs = {}

            for d in self.lifterData:

                if d['Day'] == i:
                    lifterTemplate = json.load(open(self.input.lifterJSONTemplate))

                    lifterTemplate["lot"] = d["Lot"]
                    lifterTemplate["guest"] = False
                    lifterTemplate["novice"] = False
                    lifterTemplate["name"] = f'{d["First name"]} {d["Last name"]}'
                    lifterTemplate["divisions"][0] = d["Age"]
                    lifterTemplate["flight"] = d["Flight"].upper()
                    lifterTemplate["day"] = 1
                    lifterTemplate["instagram"] = d["Instagram"]
                    lifterTemplate["team"] = d["Association"]
                    lifterTemplate["intendedWeightClassKg"] = d["Weight"]
                    lifterTemplate["notes"] = d["Notes"]
                    lifterTemplate["country"] = "New Zealand"
                    lifterTemplate["id"] = id
                    lifterTemplate["birthDate"] = d["DOB"].strftime("%Y-%m-%d")
                    lifterTemplate["age"] = self.eventDate.year - d["DOB"].year - ((self.eventDate.month, self.eventDate.day) < (d["DOB"].month, d["DOB"].day))
                    lifterTemplate["events"][0] = "SBD"

                    sex = d["Gender"]
                    if sex.lower() == "male":
                        lifterTemplate["sex"] = "M"
                    else:
                        lifterTemplate["sex"] = "F"

                    equipment = d["Raw"]
                    if equipment == "Raw":
                        lifterTemplate["equipment"] = "Sleeves"
                    else:
                        lifterTemplate["equipment"] = "Wraps"

                    lifterTemplate["canBreakRecords"] = True
                    
                    outputs[id] = lifterTemplate

                    id += 1
            self.log(f"OL Data Day {i}")

            OLTemplate = json.load(open(self.input.olTemplate))

            OLTemplate["meet"]["name"] = f"{self.input.eventName} Session {i}"

            dayAdv = (i - 1) // 2
            date = self.eventDate + datetime.timedelta(days=dayAdv)
            OLTemplate["meet"]["date"] = date.strftime("%Y-%m-%d")
            OLTemplate["meet"]["lengthDays"] = 1
            OLTemplate["meet"]["platformsOnDays"] = [1]

            OLTemplate["registration"]["entries"] = list(outputs.values())

            lookup = {}
            for k in list(outputs.keys()):
                lookup[int(k)]= int(k)

            OLTemplate["registration"]["lookup"] = lookup
            
            with open(f"{self.input.eventName} Session {i}.openlifter", 'w', newline='') as file:
                file.write(json.dumps(OLTemplate).replace("'", ""))

    def checkInputs(self):
        r = True
        self.input.eventName = self.eventNameField.get()
        for attr, value in vars(self.input).items():
            if value is not None:
                pass
            else:
                self.log(f"Please populate {attr}")
                r = False

        return r
    
    def run(self):
        tmp = tempfile.TemporaryDirectory()
        self.temp = tmp.name

        if self.checkInputs():
            self.proccessData()
            #self.GUID = str(uuid.uuid4())
            #os.mkdir(os.path.join(self.temp, self.GUID))


            self.createOLData()
            self.runLifterSpecific()
            self.createWeighIn()
            self.createGearCheck()
            self.createManualScoreSheet()
            # self.cleanUp()


            self.log("~~~~~~~~~~~~~~~~~\nCompleted!\n~~~~~~~~~~~~~~~~~")
        

class PDFGenerator:
    def __init__(self):
        self.speakerPDFs = []
        self.certificatePDFs= []
    def createSpeaker(self, data):
        lifter, speakerTemplate, temp = data
        # Create a new Word docx.document based on the provided template
        doc = docx.Document(speakerTemplate)


        name = f'{lifter["First name"]} {lifter["Last name"]}'
        dob = lifter["DOB"].strftime("%d/%m/%Y")
        
        #spaces are so it doesn't replace other numbers in the template
        self.findReplaceTable(doc,dict(zzNAMEzz=name,zzCLASSzz=lifter["Weight"], zzDOBzz=dob, zzNATIONzz=lifter["Nation"], zzLOTzz=lifter["Lot"]+"   "))

        time.sleep(0.1)
        path = os.path.join(temp, f'speaker_{name}.docx')
        doc.save(path)
        time.sleep(0.1)
        return(path)

    def createCetificates(self, data):

        lifter, certificateTemplate, temp = data

        doc = docx.Document(certificateTemplate)

        name = f'{lifter["First name"]} {lifter["Last name"]}'

        self.findReplaceTable(doc,dict(zzNAMEzz=name,zzCLASSzz=lifter["Weight"], zzGENDERzz=lifter["Gender"], zzDIVzz=lifter["Age"], zzEQUIPzz=lifter["Raw"]))

        time.sleep(0.1)
        path = os.path.join(temp, f'certificate_{name}.docx')
        doc.save(path)
        time.sleep(0.1)
        return (path)

    def findReplaceTable(self, doc, data, revert=False):
        """
        Find and replace text in tables in a Word docx.document.

        Arguments:
        - doc: a `docx.Document` object representing the Word document to modify
        - data: a dictionary containing the text to find and its replacement
        - revert: if `True`, the find and replace arguments will be swapped
            (default: False)

        Returns:
        - None

        """
        # If revert is True, swap find and replace arguments
        if revert:
            data = {v: k for k, v in data.items()}

        # Loop over the tables in the document
        for table in doc.tables:
            # Loop over the rows in the table
            for row in table.rows:
                # Loop over the cells in the row
                for cell in row.cells:
                    # If the find text is in the cell, loop over the paragraphs in the cell
                    if any(find in cell.text for find in data.keys()):
                        for p in cell.paragraphs:
                            # If the find text is in the paragraph, loop over the runs in the paragraph
                            if any(find in p.text for find in data.keys()):
                                for run in p.runs:
                                    # If the find text is in the run, replace it with the replacement text
                                    if any(find in run.text for find in data.keys()):
                                        for find, replace in data.items():
                                            run.text = run.text.replace(find, replace)
        




if __name__ == '__main__':
    multiprocessing.freeze_support()
    app = CPA_Cert_Generator()
    app.mainloop()
    print()